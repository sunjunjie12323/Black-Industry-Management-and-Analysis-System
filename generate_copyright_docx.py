import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = r"c:\Users\sunjunjie\Desktop\vibe coding项目\黑产系统\threat-intel-agent"
SRC_BACKEND = os.path.join(BASE_DIR, "backend", "app")
SRC_FRONTEND = os.path.join(BASE_DIR, "frontend", "src")
OUTPUT_DIR = os.path.join(BASE_DIR, "软著源码")

SOFTWARE_LIST = [
    {
        "id": 1,
        "name": "黑灰产情报自动化采集与处理系统",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/intelligence.py"),
            (SRC_BACKEND, "api/intel_pipeline.py"),
            (SRC_BACKEND, "api/alerts.py"),
            (SRC_BACKEND, "api/blacktalk.py"),
            (SRC_BACKEND, "api/reports.py"),
            (SRC_BACKEND, "api/pirs.py"),
            (SRC_BACKEND, "api/entities.py"),
            (SRC_BACKEND, "core/auto_collector.py"),
            (SRC_BACKEND, "core/intel_pipeline.py"),
            (SRC_BACKEND, "core/alert_engine.py"),
            (SRC_BACKEND, "core/blacktalk_engine.py"),
            (SRC_BACKEND, "core/source_scheduler.py"),
            (SRC_BACKEND, "core/stix_exporter.py"),
            (SRC_BACKEND, "core/pir_engine.py"),
            (SRC_BACKEND, "core/report_generator.py"),
            (SRC_BACKEND, "core/rule_based_extractor.py"),
            (SRC_BACKEND, "core/case_manager.py"),
            (SRC_BACKEND, "core/ner_engine.py"),
            (SRC_FRONTEND, "pages/Intelligence.tsx"),
            (SRC_FRONTEND, "pages/AlertCenter.tsx"),
            (SRC_FRONTEND, "pages/Entities.tsx"),
            (SRC_FRONTEND, "pages/PIRs.tsx"),
            (SRC_FRONTEND, "pages/IntelPipeline.tsx"),
            (SRC_FRONTEND, "pages/Reports.tsx"),
        ],
    },
    {
        "id": 2,
        "name": "基于大模型的黑灰产威胁智能分析引擎",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/deepseek.py"),
            (SRC_BACKEND, "api/threat_analysis.py"),
            (SRC_BACKEND, "api/attack_prediction.py"),
            (SRC_BACKEND, "api/zero_day.py"),
            (SRC_BACKEND, "api/attribution.py"),
            (SRC_BACKEND, "api/provenance.py"),
            (SRC_BACKEND, "api/temporal_decay.py"),
            (SRC_BACKEND, "core/attack_chain_predictor.py"),
            (SRC_BACKEND, "core/zero_day_detector.py"),
            (SRC_BACKEND, "core/entity_attribution.py"),
            (SRC_BACKEND, "core/provenance_chain.py"),
            (SRC_BACKEND, "core/temporal_decay.py"),
            (SRC_BACKEND, "core/evidence_chain.py"),
            (SRC_BACKEND, "core/llm.py"),
            (SRC_BACKEND, "core/local_embedding.py"),
            (SRC_BACKEND, "core/vector_store.py"),
            (SRC_BACKEND, "engine/analysis_engine.py"),
            (SRC_BACKEND, "engine/deep_analysis.py"),
            (SRC_BACKEND, "engine/analyzers/attack_prediction.py"),
            (SRC_BACKEND, "engine/analyzers/zero_day.py"),
            (SRC_BACKEND, "engine/analyzers/attribution.py"),
            (SRC_BACKEND, "engine/analyzers/provenance.py"),
            (SRC_BACKEND, "engine/analyzers/decay.py"),
            (SRC_FRONTEND, "pages/AttackPrediction.tsx"),
            (SRC_FRONTEND, "pages/ZeroDayDetection.tsx"),
            (SRC_FRONTEND, "pages/Attribution.tsx"),
            (SRC_FRONTEND, "pages/Provenance.tsx"),
            (SRC_FRONTEND, "pages/DecayTracking.tsx"),
        ],
    },
    {
        "id": 3,
        "name": "黑灰产知识图谱构建与关联分析系统",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/graph.py"),
            (SRC_BACKEND, "api/entities.py"),
            (SRC_BACKEND, "api/ner.py"),
            (SRC_BACKEND, "core/knowledge_graph.py"),
            (SRC_BACKEND, "core/ner_engine.py"),
            (SRC_BACKEND, "core/entity_attribution.py"),
            (SRC_BACKEND, "core/rule_based_extractor.py"),
            (SRC_BACKEND, "core/vector_store.py"),
            (SRC_BACKEND, "core/local_embedding.py"),
            (SRC_FRONTEND, "pages/Graph.tsx"),
            (SRC_FRONTEND, "pages/Entities.tsx"),
        ],
    },
    {
        "id": 4,
        "name": "面向黑灰产领域的提示词工程与模型微调平台",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/prompt_engine.py"),
            (SRC_BACKEND, "api/data_pipeline.py"),
            (SRC_BACKEND, "api/finetune.py"),
            (SRC_BACKEND, "api/domain_finetune.py"),
            (SRC_BACKEND, "core/prompt_engine_service.py"),
            (SRC_BACKEND, "core/pipeline_engine.py"),
            (SRC_BACKEND, "core/finetune_engine.py"),
            (SRC_BACKEND, "core/domain_finetune.py"),
            (SRC_BACKEND, "core/validators.py"),
            (SRC_BACKEND, "core/llm.py"),
            (SRC_FRONTEND, "pages/PromptEngine.tsx"),
            (SRC_FRONTEND, "pages/DataPipeline.tsx"),
            (SRC_FRONTEND, "pages/ModelFinetune.tsx"),
        ],
    },
    {
        "id": 5,
        "name": "黑灰产情报智能问答与内容生成系统",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/smartqa.py"),
            (SRC_BACKEND, "api/translation.py"),
            (SRC_BACKEND, "api/content_gen.py"),
            (SRC_BACKEND, "api/industry_scene.py"),
            (SRC_BACKEND, "core/qa_engine.py"),
            (SRC_BACKEND, "core/translation_engine.py"),
            (SRC_BACKEND, "core/content_engine.py"),
            (SRC_BACKEND, "core/vector_store.py"),
            (SRC_BACKEND, "core/local_embedding.py"),
            (SRC_BACKEND, "core/llm.py"),
            (SRC_FRONTEND, "pages/SmartQA.tsx"),
            (SRC_FRONTEND, "pages/Translation.tsx"),
            (SRC_FRONTEND, "pages/ContentGeneration.tsx"),
            (SRC_FRONTEND, "pages/IndustryScene.tsx"),
        ],
    },
    {
        "id": 6,
        "name": "黑灰产态势感知与数据分析看板系统",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/dashboard.py"),
            (SRC_BACKEND, "api/data_analytics.py"),
            (SRC_BACKEND, "api/economic.py"),
            (SRC_BACKEND, "api/innovation.py"),
            (SRC_BACKEND, "core/analytics_engine.py"),
            (SRC_BACKEND, "core/economic_engine.py"),
            (SRC_BACKEND, "core/economic_integration.py"),
            (SRC_BACKEND, "core/real_data_provider.py"),
            (SRC_BACKEND, "core/economic_organism_bridge.py"),
            (SRC_FRONTEND, "pages/Dashboard.tsx"),
            (SRC_FRONTEND, "pages/DataAnalytics.tsx"),
        ],
    },
    {
        "id": 7,
        "name": "黑灰产情报安全合规与数据治理系统",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/compliance.py"),
            (SRC_BACKEND, "api/audit_log.py"),
            (SRC_BACKEND, "api/dfx.py"),
            (SRC_BACKEND, "api/backup.py"),
            (SRC_BACKEND, "api/billing.py"),
            (SRC_BACKEND, "core/data_masking.py"),
            (SRC_BACKEND, "core/data_governance.py"),
            (SRC_BACKEND, "core/audit_middleware.py"),
            (SRC_BACKEND, "core/dfx.py"),
            (SRC_BACKEND, "core/backup.py"),
            (SRC_BACKEND, "core/billing.py"),
            (SRC_FRONTEND, "pages/AuditLog.tsx"),
            (SRC_FRONTEND, "pages/Settings.tsx"),
        ],
    },
    {
        "id": 8,
        "name": "黑灰产情报系统安全认证与访问控制平台",
        "version": "V1.0",
        "files": [
            (SRC_BACKEND, "api/auth.py"),
            (SRC_BACKEND, "api/deployment.py"),
            (SRC_BACKEND, "core/auth.py"),
            (SRC_BACKEND, "core/rate_limiter.py"),
            (SRC_BACKEND, "core/api_key_manager.py"),
            (SRC_BACKEND, "core/api_key_auth.py"),
            (SRC_BACKEND, "core/tenant.py"),
            (SRC_BACKEND, "core/tenant_manager.py"),
            (SRC_BACKEND, "core/tenant_middleware.py"),
            (SRC_BACKEND, "core/cache_service.py"),
            (SRC_BACKEND, "core/metrics.py"),
            (SRC_BACKEND, "core/metrics_middleware.py"),
            (SRC_FRONTEND, "pages/Login.tsx"),
            (SRC_FRONTEND, "pages/DeploymentManage.tsx"),
        ],
    },
]

LINES_PER_PAGE = 50
TOTAL_PAGES = 60


def clean_line(line: str) -> str:
    line = re.sub(r'#.*?(作者|author|版权|copyright|地址|address|时间|created|date|@\w+)', '#', line)
    line = re.sub(r'//.*?(作者|author|版权|copyright|地址|address|时间|created|date|@\w+)', '//', line)
    line = re.sub(r'/\*.*?(作者|author|版权|copyright|地址|address|时间|created|date|@\w+).*?\*/', '', line)
    line = re.sub(r'\{\*.*?\*\}', '', line)
    return line


def is_blank(line: str) -> bool:
    return line.strip() == ''


def process_file(src_path: str) -> list[str]:
    if not os.path.exists(src_path):
        return []
    with open(src_path, 'r', encoding='utf-8', errors='replace') as f:
        raw_lines = f.readlines()
    result = []
    for line in raw_lines:
        cleaned = clean_line(line.rstrip('\n'))
        if is_blank(cleaned):
            continue
        result.append(cleaned)
    return result


def add_header(doc, software_name, version, page_num, total_pages):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(f"{software_name} {version} 源代码 - 第{page_num}页 共{total_pages}页")
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    pPr = header_para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(software: dict) -> None:
    name = software["name"]
    version = software["version"]
    folder_name = f"{software['id']:02d}_{name}"
    folder_path = os.path.join(OUTPUT_DIR, folder_name)
    os.makedirs(folder_path, exist_ok=True)

    all_lines = []
    file_sections = []

    for base, rel_path in software["files"]:
        full_path = os.path.join(base, rel_path)
        lines = process_file(full_path)
        if lines:
            ext = os.path.splitext(rel_path)[1]
            if ext == '.py':
                header = f"# {'=' * 60}"
                sep = f"# File: {rel_path}"
                footer = f"# {'=' * 60}"
            else:
                header = f"// {'=' * 60}"
                sep = f"// File: {rel_path}"
                footer = f"// {'=' * 60}"
            section = [header, sep, footer] + lines
            file_sections.append((rel_path, section))
            all_lines.extend(section)

    total_lines = len(all_lines)

    if total_lines <= TOTAL_PAGES * LINES_PER_PAGE:
        selected_lines = all_lines
        actual_pages = (len(selected_lines) + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    else:
        front_count = 30 * LINES_PER_PAGE
        front_lines = all_lines[:front_count]
        back_start = total_lines - 30 * LINES_PER_PAGE
        back_lines = all_lines[back_start:]
        selected_lines = front_lines + back_lines
        actual_pages = TOTAL_PAGES

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = Pt(15)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.page_break_before = False

    for page_idx in range(actual_pages):
        start = page_idx * LINES_PER_PAGE
        end = start + LINES_PER_PAGE
        page_lines = selected_lines[start:end]

        if not page_lines:
            break

        if page_idx > 0:
            doc.add_page_break()

        page_text = '\n'.join(page_lines)
        para = doc.add_paragraph(page_text)
        para.paragraph_format.line_spacing = Pt(15)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        for run in para.runs:
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    from docx.oxml import parse_xml
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run1 = header_para.add_run(f"{name} {version} 源代码 - 第")
    run1.font.name = '宋体'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run1.font.size = Pt(9)

    run2 = header_para.add_run()
    run2._element.append(fldChar1)
    run2._element.append(instrText)
    run2._element.append(fldChar2)
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(9)

    run3 = header_para.add_run(f"页 共{actual_pages}页")
    run3.font.name = '宋体'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run3.font.size = Pt(9)

    pPr = header_para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    section.footer.is_linked_to_previous = False

    out_file = os.path.join(folder_path, f"{name} {version}_源代码.docx")
    doc.save(out_file)

    info_file = os.path.join(folder_path, "源代码说明.txt")
    with open(info_file, 'w', encoding='utf-8') as f:
        f.write(f"软件名称：{name}\n")
        f.write(f"版本号：{version}\n")
        f.write(f"总代码行数（不含空行）：{total_lines}\n")
        f.write(f"提交页数：{actual_pages}页\n")
        f.write(f"每页行数：{LINES_PER_PAGE}行\n")
        f.write(f"包含源文件：\n")
        for rel_path, section in file_sections:
            f.write(f"  - {rel_path} ({len(section) - 3} 行)\n")
        f.write(f"\n格式说明：\n")
        f.write(f"  - Word文档(.docx)格式\n")
        f.write(f"  - 字体：Consolas 10号\n")
        f.write(f"  - 行距：固定15磅\n")
        f.write(f"  - 页边距：上下左右2.5cm\n")
        f.write(f"  - 页眉：软件名称+版本号+页码\n")
        f.write(f"  - 已删除含人名/地址/时间/版权的注释\n")
        f.write(f"  - 每页50行有效代码（不含空行）\n")
        if total_lines > TOTAL_PAGES * LINES_PER_PAGE:
            f.write(f"  - 代码超过3000行，提交前30页+后30页\n")
        else:
            f.write(f"  - 代码不足3000行，提交全部代码\n")

    print(f"  [{software['id']}] {name}: {total_lines}行 -> {actual_pages}页 -> {out_file}")


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"输出目录: {OUTPUT_DIR}\n")
    print("=" * 60)
    print("软著源代码Word文档生成工具")
    print("格式规范：Consolas 10号 / 行距15磅 / 页边距2.5cm / 每页50行 / 页眉含软件名称+页码")
    print("=" * 60)

    for sw in SOFTWARE_LIST:
        print(f"\n处理: [{sw['id']}] {sw['name']} {sw['version']}")
        generate_docx(sw)

    print("\n" + "=" * 60)
    print("全部完成! 请打开Word文档检查格式。")
    print("=" * 60)


if __name__ == "__main__":
    main()
